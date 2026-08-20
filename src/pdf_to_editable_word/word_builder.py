from __future__ import annotations

from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from .document_model import DocumentModel, ImageObject, PageModel, VectorObject, TextSpan
from .font_resolver import FontResolver

def _rgb(value: int) -> str:
    return f"{value & 0xFFFFFF:06X}"


def _frame_width_twips(width_points: float) -> int:
    """Allow for small metric differences between embedded PDF fonts and Word fonts."""
    return max(round((width_points + 12) * 20), 20)


class PositionedWordBuilder:
    """Writes editable, page-positioned Word text and image objects."""

    def build(self, model: DocumentModel, output_path: Path) -> None:
        document = Document()
        self.font_resolver = FontResolver()
        with TemporaryDirectory(prefix="pdf2word-media-") as media_dir:
            for index, page in enumerate(model.pages):
                section = document.sections[0] if index == 0 else document.add_section(WD_SECTION.NEW_PAGE)
                self._configure_section(section, page)
                self._add_page_objects(document, page, Path(media_dir), index * 100000)
        document.save(output_path)

    @staticmethod
    def _configure_section(section, page: PageModel) -> None:
        section.page_width = Pt(page.width)
        section.page_height = Pt(page.height)
        section.top_margin = Pt(0)
        section.bottom_margin = Pt(0)
        section.left_margin = Pt(0)
        section.right_margin = Pt(0)
        section.header_distance = Pt(0)
        section.footer_distance = Pt(0)

    def _add_page_objects(self, document: Document, page: PageModel, media_dir: Path, base_id: int) -> None:
        # A scanned page image is intentionally underneath the recognized text layer.
        for index, image in enumerate(image for image in page.images if image.is_page_background):
            self._add_image(document, image, media_dir, base_id + 70000 + index, z_index=0)
        for index, image in enumerate(
            image
            for image in page.images
            if not image.is_page_background and not image.is_stamp
        ):
            self._add_image(document, image, media_dir, base_id + 80000 + index, z_index=20)
        for index, vector in enumerate(page.vectors):
            self._add_vector(document, vector, media_dir, base_id + 90000 + index, z_index=30)
        for index, span in enumerate(page.text_spans):
            self._add_textbox(document, span, base_id + index + 1, z_index=40)
        for index, image in enumerate(image for image in page.images if image.is_stamp):
            self._add_image(document, image, media_dir, base_id + 50000 + index, z_index=50)
        document.add_paragraph()

    def _add_textbox(self, document: Document, span: TextSpan, shape_id: int, z_index: int) -> None:
        # Frame paragraphs are editable and preserve a PDF text span's page coordinates.
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph_properties = paragraph._p.get_or_add_pPr()
        frame = OxmlElement("w:framePr")
        frame.set(qn("w:w"), str(_frame_width_twips(span.bbox.width)))
        frame.set(qn("w:h"), str(max(round(span.bbox.height * 20), 20)))
        frame.set(qn("w:x"), str(round(span.bbox.x0 * 20)))
        # PDF glyph boxes begin at the ascender while Word frames begin at paragraph layout top.
        # OCR boxes already follow the bitmap glyph top, so they need a separate correction.
        top = span.bbox.y0 - 3.0 if span.source == "ocr" else span.bbox.y0 + 2.5
        frame.set(qn("w:y"), str(round(top * 20)))
        frame.set(qn("w:hAnchor"), "page")
        frame.set(qn("w:vAnchor"), "page")
        frame.set(qn("w:wrap"), "notBeside")
        frame.set(qn("w:hSpace"), "0")
        frame.set(qn("w:vSpace"), "0")
        frame.set(qn("w:zIndex"), str(z_index))
        paragraph_properties.append(frame)
        run = paragraph.add_run()
        properties = run._r.get_or_add_rPr()
        fonts = OxmlElement("w:rFonts")
        font_name = self.font_resolver.resolve(span.font_name)
        for attribute in ("ascii", "hAnsi", "eastAsia"):
            fonts.set(qn(f"w:{attribute}"), font_name)
        properties.append(fonts)
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), str(max(round(span.font_size * 2), 8)))
        properties.append(size)
        if span.bold:
            properties.append(OxmlElement("w:b"))
        if span.italic:
            properties.append(OxmlElement("w:i"))
        if span.underline:
            underline = OxmlElement("w:u")
            underline.set(qn("w:val"), "single")
            properties.append(underline)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), _rgb(span.color))
        properties.append(color)
        # Compress only when Word's substituted font is wider than the original PDF font.
        # This prevents an otherwise single-line PDF span from reflowing onto a second line.
        fit_text = OxmlElement("w:fitText")
        fit_text.set(qn("w:val"), str(_frame_width_twips(span.bbox.width)))
        properties.append(fit_text)
        run.add_text(span.text)

    def _add_image(
        self, document: Document, image: ImageObject, media_dir: Path, shape_id: int, z_index: int
    ) -> None:
        suffix = ".png" if image.is_stamp else f".{image.extension.lstrip('.')}"
        image_path = media_dir / f"image-{shape_id}{suffix}"
        image_path.write_bytes(image.data)
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph_properties = paragraph._p.get_or_add_pPr()
        frame = OxmlElement("w:framePr")
        frame.set(qn("w:w"), str(_frame_width_twips(image.bbox.width)))
        frame.set(qn("w:h"), str(max(round(image.bbox.height * 20), 20)))
        # Word/LibreOffice reserves a leading inline-image inset inside a frame.
        frame.set(qn("w:x"), str(round((image.bbox.x0 - 18) * 20)))
        frame.set(qn("w:y"), str(round(image.bbox.y0 * 20)))
        frame.set(qn("w:hAnchor"), "page")
        frame.set(qn("w:vAnchor"), "page")
        frame.set(qn("w:wrap"), "notBeside")
        frame.set(qn("w:hSpace"), "0")
        frame.set(qn("w:vSpace"), "0")
        frame.set(qn("w:zIndex"), str(z_index))
        paragraph_properties.append(frame)
        paragraph.add_run().add_picture(
            str(image_path), width=Pt(image.bbox.width), height=Pt(image.bbox.height)
        )

    def _add_vector(
        self, document: Document, vector: VectorObject, media_dir: Path, shape_id: int, z_index: int
    ) -> None:
        self._add_image(
            document,
            ImageObject(bbox=vector.bbox, data=vector.data, extension="png"),
            media_dir,
            shape_id,
            z_index,
        )
